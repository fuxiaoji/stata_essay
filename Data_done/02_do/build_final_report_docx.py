# -*- coding: utf-8 -*-
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(os.environ.get("DATA_DONE_ROOT", "/Users/Zhuanz1/Desktop/code/stata_essay/Data_done"))
REPORT = ROOT / "00_report" / "final_report.docx"
FIG_DIR = ROOT / "04_figures"
FIGURES = [
    FIG_DIR / "fig1_educ_distribution.png",
    FIG_DIR / "fig2_mean_lnwage_by_edu.png",
    FIG_DIR / "fig3_educ_lnwage_fit.png",
]
DOC_FIGURES = [
    FIG_DIR / "report_doc_fig1_educ_distribution.jpg",
    FIG_DIR / "report_doc_fig2_mean_lnwage_by_edu.jpg",
    FIG_DIR / "report_doc_fig3_educ_lnwage_fit.jpg",
]


def set_east_asia_font(target, font_name):
    target.font.name = font_name
    rpr = target.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def set_run_style(run, font_name="宋体", size=10.5, bold=False, color=None):
    set_east_asia_font(run, font_name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, font_name, size, bold=False, color=None):
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    style.font.size = Pt(size)
    style.font.bold = bold
    if color:
        style.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_paragraph(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_style(run, bold=bold)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_style(run, font_name="微软雅黑", size=15, bold=True, color="1F4E79")
    else:
        set_run_style(run, font_name="微软雅黑", size=12.5, bold=True, color="1F4E79")
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, "DCE6F1")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_style(run, font_name="微软雅黑", size=9.5, bold=True, color="1F1F1F")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_style(run, size=9.2)

    doc.add_paragraph()


def add_figure(doc, image_path, caption):
    if not image_path.exists():
        raise FileNotFoundError(f"Missing figure: {image_path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run = cap.add_run(caption)
    set_run_style(run, size=9, color="555555")


def prepare_figures():
    for source, target in zip(FIGURES, DOC_FIGURES):
        if not source.exists():
            raise FileNotFoundError(f"Missing figure: {source}")
        image = Image.open(source).convert("RGB")
        image.save(target, format="JPEG", quality=95, optimize=True, dpi=(300, 300))


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    set_style_font(doc.styles["Normal"], "宋体", 10.5)
    set_style_font(doc.styles["Heading 1"], "微软雅黑", 15, bold=True, color="1F4E79")
    set_style_font(doc.styles["Heading 2"], "微软雅黑", 12.5, bold=True, color="1F4E79")


def build_report():
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    prepare_figures()
    doc = Document()
    configure_document(doc)

    doc.core_properties.title = "CFPS 2022 教育水平与收入实证报告"
    doc.core_properties.subject = "教育年限与收入相关关系分析"
    doc.core_properties.author = "Codex"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("教育水平是否提高中国成年人的收入？")
    set_run_style(run, font_name="微软雅黑", size=18, bold=True, color="17365D")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    run = subtitle.add_run("基于 CFPS 2022 数据的实证分析")
    set_run_style(run, font_name="微软雅黑", size=12, color="666666")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run("样本量 N = 4,699；因变量为 ln(wage + 1)；标准误为稳健标准误")
    set_run_style(run, size=9.5, color="666666")

    add_heading(doc, "摘要", 1)
    add_paragraph(
        doc,
        "本文使用 CFPS 2022 数据考察教育水平与收入之间的关系。样本限定为 2022 年 18-64 岁成年人，并要求工资变量与核心控制变量可用，最终得到 4,699 个观测值。基准回归以 ln(wage + 1) 为因变量，以受教育年限为核心解释变量，并逐步加入年龄、年龄平方、性别、婚姻、城乡户口、医保和省份固定效应。结果显示，主模型中教育年限系数为 0.0882，表示在控制其他变量后，受教育年限每增加 1 年，ln(wage + 1) 平均增加约 0.0882 个对数点，近似对应工资收入提高 8.8%。稳健性检验从教育变量定义、收入变量定义和样本筛选口径三个角度展开，结论方向保持一致。由于本文使用横截面数据和常规 OLS 回归，结果应理解为教育水平与收入之间的条件相关关系，而非严格因果效应。"
    )

    add_heading(doc, "一、引言", 1)
    add_paragraph(
        doc,
        "教育通常被视为提升个体劳动市场表现的重要渠道，更高的教育水平可能通过提高技能、改善职业匹配和扩展就业机会影响收入。本文关注的问题是：在 CFPS 2022 的中国成年样本中，教育水平是否与更高收入显著相关。"
    )
    add_paragraph(
        doc,
        "需要强调的是，本文使用横截面数据和常规 OLS 回归，结果主要反映教育与收入之间的条件相关关系，不能直接解释为严格因果效应。个体能力、家庭背景、学校质量、行业与职业选择等因素仍可能同时影响教育水平和收入，从而带来遗漏变量偏误。"
    )

    add_heading(doc, "二、数据来源与样本筛选", 1)
    add_paragraph(
        doc,
        "本文使用 CFPS 2022 数据构建分析样本。样本筛选口径为：保留 2022 年样本；年龄限制在 18-64 岁；工资变量 wage 非缺失；并要求 lnwage、educ、age_、age2、gen、mar、rural、medsure_dum 和 provcd 等变量可用。按此口径处理后，最终样本量为 4,699。"
    )
    add_paragraph(
        doc,
        "样本中工资收入均值为 60,594.13 元，受教育年限均值为 11.01 年，平均年龄为 40.87 岁。教育分类上，初中、高中/中专、大专和本科构成了主要样本群体。"
    )

    add_heading(doc, "三、变量定义", 1)
    add_table(
        doc,
        ["变量", "含义", "处理方式"],
        [
            ["wage", "工资收入", "主收入变量，用于生成 lnwage = ln(wage + 1)"],
            ["lnwage", "工资收入对数", "主因变量"],
            ["inc1", "家庭人均纯收入", "替代收入变量，用于生成 lninc1 = ln(inc1 + 1)"],
            ["lninc1", "家庭人均纯收入对数", "稳健性检验中的替代因变量"],
            ["educ", "受教育年限", "核心解释变量，按教育层级折算为年限"],
            ["edu", "教育分类", "稳健性检验中使用的分类变量"],
            ["age_", "年龄", "受访者年龄，限制在 18-64 岁"],
            ["age2", "年龄平方", "用于控制生命周期非线性效应"],
            ["gen", "性别", "1 = 男性，0 = 女性"],
            ["mar", "婚姻状态", "1 = 有配偶，0 = 无配偶"],
            ["rural", "城乡户口", "1 = 农业户口，0 = 非农业户口"],
            ["medsure_dum", "医保", "是否购买医保"],
            ["provcd", "省份代码", "用于加入省份固定效应"],
        ],
        [1.3, 2.1, 3.0],
    )

    add_heading(doc, "四、模型设定", 1)
    add_paragraph(doc, "基准模型可写为：")
    add_paragraph(
        doc,
        "ln(wage_i + 1) = α + β educ_i + γ X_i + μ_p + ε_i",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "其中，educ_i 表示个体受教育年限，X_i 包括年龄、年龄平方、性别、婚姻、城乡户口和医保等控制变量，μ_p 表示省份固定效应。加入省份固定效应后，模型能够控制不同省份之间相对稳定的地区发展水平、劳动力市场结构、工资水平和公共资源差异，从而减少地区层面差异对教育收入关系估计的干扰。"
    )
    add_paragraph(
        doc,
        "本文关注教育年限系数 β 的符号、大小和显著性，所有回归均使用 robust 稳健标准误。主模型中，教育年限系数为 0.0882，可解释为受教育年限每增加 1 年，ln(wage + 1) 平均增加约 0.0882 个对数点，近似对应工资收入提高 8.8%。"
    )

    add_heading(doc, "五、描述性统计与图形证据", 1)
    add_table(
        doc,
        ["变量", "样本量", "均值", "标准差", "最小值", "最大值"],
        [
            ["wage", "4,699", "60,594.13", "68,985.46", "0", "1,800,000"],
            ["lnwage", "4,699", "10.4892", "1.5671", "0", "14.4033"],
            ["educ", "4,699", "11.0062", "4.3765", "0", "22"],
            ["age_", "4,699", "40.8668", "10.9207", "18", "64"],
            ["age2", "4,699", "1,789.3310", "919.7751", "324", "4,096"],
        ],
        [1.1, 1.0, 1.2, 1.2, 1.0, 1.0],
    )
    add_paragraph(
        doc,
        "描述性统计表明，样本受教育年限中位数约为 12 年。工资收入分布存在较强右偏，因此本文使用 ln(wage + 1) 作为主因变量，以降低极端值对估计结果的影响。"
    )
    add_figure(doc, DOC_FIGURES[0], "图 1 教育年限分布。数据来源：CFPS 2022。")
    add_figure(doc, DOC_FIGURES[1], "图 2 不同教育水平的平均 ln(wage + 1)。")
    add_figure(doc, DOC_FIGURES[2], "图 3 教育年限与收入之间的拟合关系。")
    add_paragraph(
        doc,
        "三张图从不同角度展示了教育与收入的正向关系：图 1 给出样本的教育结构，图 2 显示更高教育层级通常对应更高的平均对数工资，图 3 则表明教育年限组均值与收入之间存在明显的正向拟合关系。"
    )

    add_heading(doc, "六、基准回归结果", 1)
    add_table(
        doc,
        ["变量", "模型(1)", "模型(2)", "模型(3)"],
        [
            ["educ", "0.0996***\n(0.0052)", "0.0959***\n(0.0063)", "0.0882***\n(0.0062)"],
            ["age_", "-", "0.1336***\n(0.0201)", "0.1379***\n(0.0197)"],
            ["age2", "-", "-0.0017***\n(0.0002)", "-0.0018***\n(0.0002)"],
            ["gen", "-", "0.5908***\n(0.0448)", "0.5802***\n(0.0447)"],
            ["mar", "-", "-0.0149\n(0.0514)", "0.0166\n(0.0517)"],
            ["rural", "-", "0.0482\n(0.0599)", "0.0379\n(0.0613)"],
            ["medsure_dum", "-", "-0.0177\n(0.0939)", "-0.0020\n(0.0909)"],
            ["省份固定效应", "否", "否", "是"],
            ["样本量", "4,699", "4,699", "4,699"],
            ["R²", "0.0773", "0.1272", "0.1577"],
        ],
        [1.35, 1.55, 1.55, 1.55],
    )
    add_paragraph(
        doc,
        "注：括号内为稳健标准误；*** 表示在 1% 水平显著。模型(1)仅包含教育年限，模型(2)加入人口学和家庭控制变量，模型(3)进一步加入省份固定效应。"
    )
    add_paragraph(
        doc,
        "基准回归显示，教育年限系数在三个模型中均显著为正。随着控制变量逐步加入，教育年限系数由 0.0996 下降至 0.0882，但仍在 1% 水平显著。这说明在控制年龄、性别、婚姻、城乡户口、医保和省份固定效应后，教育水平与收入之间仍存在稳定的正相关关系。"
    )
    add_paragraph(
        doc,
        "以模型(3)为主模型，教育年限系数为 0.0882，表示在控制其他变量后，受教育年限每增加 1 年，ln(wage + 1) 平均增加约 0.0882 个对数点，近似对应工资收入提高 8.8%。"
    )

    add_heading(doc, "七、稳健性检验", 1)
    add_table(
        doc,
        ["检验", "设定", "核心结果", "样本量", "R²", "结论"],
        [
            ["教育分类变量", "以 i.edu 替代 educ", "较高教育层级系数均为正且显著", "4,699", "0.1670", "方向一致"],
            ["替代收入变量", "lninc1 = ln(inc1 + 1)", "educ = 0.0728***\n(0.0030)", "4,699", "0.3395", "仍显著为正"],
            ["正工资样本", "仅保留 wage > 0", "educ = 0.0788***\n(0.0037)", "4,638", "0.2970", "仍显著为正"],
        ],
        [1.15, 1.55, 1.95, 0.85, 0.7, 1.0],
    )
    add_paragraph(
        doc,
        "本文从教育变量定义、收入变量定义和样本筛选口径三个角度进行稳健性检验，以检查基准结论是否依赖于特定变量构造或样本处理方式。"
    )
    add_paragraph(
        doc,
        "第一，使用教育分类变量替代教育年限变量，可以检验结论是否依赖于教育年限的折算方式。第二，使用替代收入变量 lninc1，可以检验教育与收入的正相关关系是否只存在于工资收入指标中。第三，仅保留正工资样本，可以检验零工资样本是否对基准估计结果产生明显影响。三项检验的结果均表明，教育变量系数保持显著为正，因此基准结论较为稳健。"
    )

    add_heading(doc, "八、结论与局限", 1)
    add_paragraph(
        doc,
        "本文基于 CFPS 2022 数据分析了中国成年人教育水平与收入之间的关系。结果显示，无论是在简单模型、加入控制变量的模型，还是加入省份固定效应的模型中，教育年限系数均显著为正。主模型估计表明，受教育年限每增加 1 年，ln(wage + 1) 平均增加约 0.0882 个对数点，近似对应工资收入提高 8.8%。"
    )
    add_paragraph(
        doc,
        "同时，本文结论应理解为条件相关关系，不能直接解释为严格因果效应。除横截面数据本身的限制外，个体能力、家庭背景、学校质量、行业与职业选择等因素仍可能同时影响教育水平和收入。尽管如此，基准回归和稳健性检验均指向一致结论：教育水平与收入之间存在稳定且显著的正相关关系。"
    )

    add_heading(doc, "附录：复现包文件说明", 1)
    add_table(
        doc,
        ["文件", "用途", "是否建议提交"],
        [
            ["00_master.do", "主控脚本，按顺序调用数据整理、回归、图表和 CSV 导出", "建议"],
            ["01_build_analysis_data.do", "从原始数据构建分析样本", "建议"],
            ["02_baseline_regression.do", "运行描述统计、基准回归和稳健性检验", "建议"],
            ["03_make_report_figures.do", "生成三张报告图表", "建议"],
            ["04_export_analysis_csv.do", "导出 analysis.csv 供预览检查", "建议"],
            ["analysis.dta", "清洗后的分析数据", "建议"],
            ["analysis.csv", "分析数据文本版", "建议"],
            ["regression.log", "正式回归日志", "建议"],
            ["final_report.docx", "最终中文报告", "建议"],
        ],
        [1.7, 3.7, 1.1],
    )

    doc.save(REPORT)
    return REPORT


if __name__ == "__main__":
    path = build_report()
    print(f"saved={path}")
